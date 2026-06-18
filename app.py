import streamlit as st
import json
import re
import io
import random
import os
from openai import OpenAI


# ── Deck counter + daily free cap ─────────────────────────────────
_COUNT_FILE = "/tmp/pitchscan_count.json"
_COUNT_SEED = 47
_DAILY_FREE_LIMIT = 20

def _load_data() -> dict:
    try:
        if os.path.exists(_COUNT_FILE):
            with open(_COUNT_FILE) as f:
                return json.load(f)
    except Exception:
        pass
    return {}

def _save_data(data: dict):
    try:
        with open(_COUNT_FILE, "w") as f:
            json.dump(data, f)
    except Exception:
        pass

def get_deck_count() -> int:
    return _load_data().get("count", _COUNT_SEED)

def get_daily_used() -> int:
    import datetime
    data = _load_data()
    today = datetime.date.today().isoformat()
    if data.get("day") != today:
        return 0
    return data.get("day_count", 0)

def increment_deck_count() -> int:
    import datetime
    data = _load_data()
    today = datetime.date.today().isoformat()
    if data.get("day") != today:
        data["day"] = today
        data["day_count"] = 0
    data["day_count"] = data.get("day_count", 0) + 1
    data["count"] = data.get("count", _COUNT_SEED) + 1
    _save_data(data)
    return data["count"]

def daily_cap_reached() -> bool:
    return get_daily_used() >= _DAILY_FREE_LIMIT

def daily_remaining() -> int:
    return max(0, _DAILY_FREE_LIMIT - get_daily_used())

st.set_page_config(
    page_title="PitchScan – VC-Grade Pitch Analysis",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="collapsed",
)

NOVUS_APP_ID = "f2950f7e-4ee3-4154-81f4-c3f25610a0fc"

VC_TIPS = [
    ("Paul Graham (YC)", "The best decks tell a story, not a business plan. Lead with a specific customer who has a painful problem."),
    ("Sequoia Capital", "Answer 'Why Now?' — timing kills more startups than bad products. Show why this moment is the right moment."),
    ("a16z", "Investors fund people first, ideas second. Spend at least 20% of your deck on team and unfair advantage."),
    ("First Round Capital", "Revenue > users > engagement > TAM claims. Show real traction even if it's small."),
    ("Bill Gurley (Benchmark)", "Unit economics must work at current scale before you pour fuel on the fire. Show CAC/LTV explicitly."),
    ("YC Advice", "Keep it under 12 slides. Every slide should answer one question, clearly. Complexity signals confusion."),
    ("Sequoia Capital", "Never claim 'no competition' — it tells investors either there's no market or you haven't looked hard enough."),
    ("Mike Maples Jr.", "The best pitches have a 'Why Now' insight that feels inevitable in hindsight but non-obvious today."),
    ("David Sacks (Craft)", "A hockey stick without explaining the inflection point is a red flag. What changed? Why then?"),
    ("Elad Gil", "Net Dollar Retention above 110% is one of the strongest signals for a B2B Series A. Show cohort retention."),
    ("YC Advice", "Your ask must be specific: '$X to reach Y milestone in Z months.' Vague asks signal vague thinking."),
    ("Sequoia Capital", "Three durable moats: network effects, switching costs, unique proprietary data. Show which one is yours."),
    ("Peter Thiel", "Aim to be the last company in your market — not just better, but 10x better in a way that matters."),
    ("Sarah Tavel (Benchmark)", "The competition slide matrix where you win every column is not credible. Pick 2-3 dimensions you truly own."),
    ("Greylock Partners", "Product-market fit evidence: retention curves that flatten, strong NPS, customers who pay and refer others."),
    ("YC Advice", "Gross margin matters more than revenue. 80%+ gross margin in SaaS signals real leverage in the model."),
    ("Fred Wilson (USV)", "Show the market size bottom-up, not top-down. A slice of a massive SAM number is never convincing."),
    ("a16z", "Distribution is the new product. Show how you acquire customers cheaply and at scale — not just that you can build."),
    ("Marc Andreessen", "Software is eating the world — but only when the timing, team, and distribution align. Prove all three."),
    ("Benchmark Capital", "The best founders know exactly what they don't know. Intellectual honesty in a pitch is a feature, not a bug."),
]

CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:ital,wght@0,300;0,400;0,500;0,600;0,700;0,800;1,400&family=DM+Mono:wght@400;500&display=swap');

*{box-sizing:border-box;}
:root{
  /* ── Palette ── */
  --navy:   #0C2A38;
  --green:  #2A6B55;
  --yellow: #E8BA4A;
  --aqua:   #7ECFC8;
  --pink:   #E09590;
  --mint:   #C8E9DE;

  /* ── Derived ── */
  --bg:  #071520;
  --bg2: #0c1e2e;
  --surface: rgba(12,42,56,0.5);
  --border: rgba(126,207,200,0.1);
  --border-hi: rgba(126,207,200,0.22);
  --text: #e8f4f2;
  --muted: #5a7a82;
  --dim: #2a4050;
  --ease-spring: cubic-bezier(0.32,0.72,0,1);
  --ease-out: cubic-bezier(0.16,1,0.3,1);
}

html,body,[class*="css"]{
  font-family:"Plus Jakarta Sans",system-ui,sans-serif;
  -webkit-font-smoothing:antialiased;
}
.stApp{background:var(--bg);color:var(--text);}
section[data-testid="stSidebar"]{background:var(--bg2);}
header[data-testid="stHeader"]{display:none!important;}
footer{display:none!important;}
.block-container{max-width:1060px;padding:0 1.5rem 3rem;}

/* Noise overlay */
.stApp::before{
  content:"";position:fixed;inset:0;z-index:999;pointer-events:none;
  opacity:0.02;
  background-image:url("data:image/svg+xml,%3Csvg viewBox='0 0 200 200' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='n'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.9' numOctaves='4' stitchTiles='stitch'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23n)'/%3E%3C/svg%3E");
  background-size:140px 140px;
}

/* ── HERO ── */
.ps-hero{text-align:center;padding:3.5rem 1rem 2rem;}
.ps-hero-badge{
  display:inline-flex;align-items:center;gap:0.4rem;
  background:rgba(232,186,74,0.1);border:1px solid rgba(232,186,74,0.22);
  color:var(--yellow);font-size:0.63rem;font-weight:700;
  letter-spacing:0.18em;text-transform:uppercase;
  padding:0.3rem 0.9rem;border-radius:99px;margin-bottom:1.3rem;
}
.ps-hero-dot{width:5px;height:5px;border-radius:50%;background:var(--yellow);}
.ps-hero h1{
  font-size:2.8rem;font-weight:800;color:var(--text);
  margin-bottom:0.8rem;letter-spacing:-0.035em;line-height:1.1;
}
.ps-hero h1 span{
  background:linear-gradient(135deg,var(--yellow),#f0cc6a);
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;
}
.ps-hero p{
  color:var(--muted);font-size:1rem;max-width:52ch;margin:0 auto;
  line-height:1.75;font-weight:400;
}

/* ── GAUGE (double-bezel) ── */
.ps-gauge-wrap{
  display:flex;flex-direction:column;align-items:center;padding:1.5rem;
  background:rgba(12,42,56,0.6);border:1px solid var(--border);border-radius:20px;
  box-shadow:inset 0 1px 1px rgba(126,207,200,0.05);
}
.ps-gauge-svg{filter:drop-shadow(0 0 16px rgba(232,186,74,0.2));}
.ps-oneliner{
  color:var(--muted);font-size:0.82rem;font-style:italic;
  text-align:center;margin-top:0.8rem;line-height:1.65;padding:0 0.8rem;max-width:28ch;
}

/* Verdict */
.verdict-chip{
  display:inline-flex;align-items:center;gap:0.4rem;
  padding:0.4rem 1.1rem;border-radius:99px;
  font-size:0.75rem;font-weight:700;letter-spacing:0.06em;margin-top:1rem;
}
.verdict-INVEST {background:rgba(42,107,85,0.18);color:#5ecba1;border:1px solid rgba(42,107,85,0.4);}
.verdict-MONITOR{background:rgba(232,186,74,0.12);color:var(--yellow);border:1px solid rgba(232,186,74,0.3);}
.verdict-PASS   {background:rgba(224,149,144,0.12);color:var(--pink);border:1px solid rgba(224,149,144,0.3);}
.verdict-dot{width:5px;height:5px;border-radius:50%;background:currentColor;}

/* ── SECTION BARS ── */
.ps-bars{padding:0.4rem 0;}
.ps-bar-row{display:flex;align-items:center;gap:0.75rem;margin-bottom:0.6rem;}
.ps-bar-name{min-width:100px;font-size:0.72rem;color:var(--muted);text-transform:capitalize;font-weight:600;}
.ps-bar-bg{flex:1;height:5px;background:rgba(255,255,255,0.04);border-radius:99px;overflow:hidden;}
.ps-bar-fill{height:100%;border-radius:99px;}
.ps-bar-score{min-width:28px;text-align:right;font-size:0.7rem;font-family:"DM Mono",monospace;color:var(--muted);font-variant-numeric:tabular-nums;}

/* ── SECTION HEADINGS ── */
.ps-sh{
  font-size:0.6rem;font-weight:700;letter-spacing:0.18em;text-transform:uppercase;
  color:var(--dim);margin:2rem 0 0.75rem;padding-bottom:0.45rem;
  border-bottom:1px solid var(--border);
}

/* ── FLAG CARDS ── */
.ps-flag{
  background:rgba(224,149,144,0.05);
  border:1px solid rgba(224,149,144,0.12);
  border-left:2px solid rgba(224,149,144,0.45);
  border-radius:0 8px 8px 0;
  padding:0.65rem 1rem;margin-bottom:0.4rem;
  font-size:0.79rem;color:rgba(224,149,144,0.8);line-height:1.6;
}
/* ── SUGGESTION CARDS ── */
.ps-sug{
  background:rgba(42,107,85,0.05);
  border:1px solid rgba(42,107,85,0.15);
  border-left:2px solid rgba(42,107,85,0.45);
  border-radius:0 8px 8px 0;
  padding:0.65rem 1rem;margin-bottom:0.4rem;
  font-size:0.79rem;color:rgba(94,203,161,0.8);line-height:1.6;
}
/* ── COMPARABLE CARDS ── */
.ps-comp{
  background:rgba(12,42,56,0.5);
  border:1px solid var(--border);
  border-radius:12px;padding:0.9rem 1rem;margin-bottom:0.5rem;
  transition:border-color 0.25s var(--ease-out);
  box-shadow:inset 0 1px 1px rgba(126,207,200,0.03);
}
.ps-comp:hover{border-color:var(--border-hi);}
.ps-comp-name{color:var(--yellow);font-weight:700;font-size:0.85rem;letter-spacing:-0.01em;}
.ps-comp-desc{color:var(--muted);font-size:0.76rem;margin-top:0.25rem;line-height:1.55;}

/* ── TIPS (double-bezel) ── */
.ps-tips-box{
  background:var(--bg2);border:1px solid var(--border);
  border-radius:16px;padding:3px;margin:1.5rem 0;
}
.ps-tips-inner{
  background:rgba(12,42,56,0.8);border-radius:14px;
  padding:1.3rem 1.5rem;min-height:88px;
  box-shadow:inset 0 1px 1px rgba(126,207,200,0.04);position:relative;
}
.ps-tips-source{font-size:0.6rem;font-weight:700;letter-spacing:0.16em;text-transform:uppercase;color:var(--aqua);margin-bottom:0.4rem;}
.ps-tips-text{color:var(--muted);font-size:0.83rem;line-height:1.65;font-style:italic;}
.ps-tips-counter{position:absolute;bottom:0.7rem;right:1rem;font-size:0.6rem;color:var(--dim);font-family:"DM Mono",monospace;}

/* ── UPLOAD ── */
.ps-upload-hint{color:var(--dim);font-size:0.75rem;text-align:center;margin-top:0.75rem;}
[data-testid="stFileUploader"]{
  background:rgba(12,42,56,0.3)!important;
  border:1.5px dashed rgba(232,186,74,0.22)!important;
  border-radius:14px!important;padding:1rem!important;
  transition:border-color 0.25s!important;
}
[data-testid="stFileUploader"]:hover{border-color:rgba(232,186,74,0.45)!important;}

/* ── BUTTONS ── */
.stDownloadButton>button,.stButton>button{
  background:rgba(12,42,56,0.6)!important;
  border:1px solid var(--border)!important;
  color:var(--muted)!important;border-radius:8px!important;
  font-size:0.78rem!important;font-weight:600!important;
  padding:0.45rem 1.1rem!important;
  transition:all 0.25s var(--ease-spring)!important;
}
.stDownloadButton>button:hover,.stButton>button:hover{
  background:rgba(232,186,74,0.08)!important;
  border-color:rgba(232,186,74,0.3)!important;
  color:var(--yellow)!important;
  transform:translateY(-1px)!important;
}

/* ── SIDEBAR ── */
[data-testid="stSidebar"]{background:var(--bg2)!important;}
[data-testid="stSidebar"] .stTextInput input{
  background:rgba(12,42,56,0.6)!important;
  border:1px solid var(--border)!important;
  color:var(--text)!important;border-radius:8px!important;
}
[data-testid="stSidebar"] .stTextInput input:focus{
  border-color:rgba(232,186,74,0.35)!important;
}

hr{border-color:var(--border)!important;}
a{color:var(--aqua)!important;}
</style>
"""

TIPS_JS = """
<script>
(function(){
  var tips = TIPS_PLACEHOLDER;
  var i = 0;
  function show(){
    var t = tips[i % tips.length];
    var src = document.getElementById('ps-tip-src');
    var txt = document.getElementById('ps-tip-txt');
    var cnt = document.getElementById('ps-tip-cnt');
    if(src && txt && cnt){
      src.style.opacity=0; txt.style.opacity=0;
      setTimeout(function(){
        src.textContent = t[0];
        txt.textContent = '“' + t[1] + '”';
        cnt.textContent = (i%tips.length+1) + ' / ' + tips.length;
        src.style.transition='opacity 0.5s'; txt.style.transition='opacity 0.5s';
        src.style.opacity=1; txt.style.opacity=1;
      },300);
    }
    i++;
  }
  show();
  setInterval(show, 6000);
})();
</script>
"""


def make_tips_js():
    tips_json = json.dumps([[t[0], t[1]] for t in VC_TIPS])
    return TIPS_JS.replace("TIPS_PLACEHOLDER", tips_json)


def render_loading_tips():
    import streamlit.components.v1 as _c
    tips_json = json.dumps([[t[0], t[1]] for t in VC_TIPS])
    first = VC_TIPS[0]
    _c.html("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;600;700&family=DM+Mono:wght@400;500&display=swap');
*{box-sizing:border-box;margin:0;padding:0;}
body{background:transparent;font-family:'Plus Jakarta Sans',sans-serif;}
.tips-box{background:#0c1e2e;border:1px solid rgba(126,207,200,0.12);border-radius:16px;padding:3px;margin:2px 0;}
.tips-inner{background:rgba(12,42,56,0.85);border-radius:14px;padding:1.1rem 1.4rem 1.8rem;min-height:72px;position:relative;}
.tips-source{font-size:0.58rem;font-weight:700;letter-spacing:0.18em;text-transform:uppercase;color:#7ECFC8;margin-bottom:0.4rem;transition:opacity 0.4s ease;}
.tips-text{color:#5a7a82;font-size:0.8rem;line-height:1.6;font-style:italic;transition:opacity 0.4s ease;}
.tips-counter{position:absolute;bottom:0.5rem;right:1rem;font-size:0.58rem;color:#2a4050;font-family:'DM Mono',monospace;}
.tips-dots{position:absolute;bottom:0.6rem;left:1.2rem;display:flex;gap:5px;align-items:center;}
.dot{width:4px;height:4px;border-radius:50%;background:rgba(126,207,200,0.2);transition:background 0.3s,transform 0.3s;cursor:pointer;}
.dot.active{background:#7ECFC8;transform:scale(1.4);}
</style>
<div class="tips-box"><div class="tips-inner">
  <div class="tips-source" id="ts"></div>
  <div class="tips-text" id="tt"></div>
  <div class="tips-dots" id="td"></div>
  <div class="tips-counter" id="tc"></div>
</div></div>
<script>
var tips=__TIPS__;
var i=0,total=tips.length;
var td=document.getElementById('td');
for(var d=0;d<total;d++){
  var dot=document.createElement('span');
  dot.className='dot'+(d===0?' active':'');
  dot.setAttribute('data-i',d);
  (function(idx){dot.onclick=function(){go(idx);clearInterval(timer);timer=setInterval(next,10000);};})(d);
  td.appendChild(dot);
}
function go(idx){
  i=idx;
  var s=document.getElementById('ts'),t=document.getElementById('tt'),c=document.getElementById('tc');
  s.style.opacity=0;t.style.opacity=0;
  setTimeout(function(){
    s.textContent=tips[i][0];
    t.textContent='"'+tips[i][1]+'"';
    c.textContent=(i+1)+' / '+total;
    td.querySelectorAll('.dot').forEach(function(d,j){d.className='dot'+(j===i?' active':'');});
    s.style.opacity=1;t.style.opacity=1;
  },350);
}
function next(){go((i+1)%total);}
go(0);
var timer=setInterval(next,10000);
</script>
""".replace('__TIPS__', tips_json), height=118, scrolling=False)
def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    native_text = "\n\n".join(text_parts).strip()
    if len(native_text) >= 150:
        return native_text
    # OCR fallback
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        st.info("Image-based PDF detected — running OCR...", icon="🔍")
        images = convert_from_bytes(file_bytes, dpi=200)
        ocr_parts = []
        for img in images:
            text = pytesseract.image_to_string(img, lang="eng")
            if text.strip():
                ocr_parts.append(text.strip())
        ocr_text = "\n\n".join(ocr_parts).strip()
        return ocr_text if ocr_text else native_text
    except ImportError as e:
        st.error(f"OCR library missing: {e}")
        return native_text
    except Exception as e:
        import traceback
        st.error(f"OCR error: {type(e).__name__}: {e}")
        st.code(traceback.format_exc())
        return native_text


def analyze_pitch(pdf_text: str, api_key: str, context: dict | None = None) -> dict:
    import streamlit as _st
    _prov_cfg = _st.session_state.get("_provider", {})
    _base_url = _prov_cfg.get("base_url", "https://api.deepseek.com")
    _model    = _prov_cfg.get("model", "deepseek-chat")
    client = OpenAI(api_key=api_key, base_url=_base_url)
    system_prompt = (
        "You are a senior partner at a top-tier VC firm (a16z, Sequoia, YC). "
        "Evaluate startup pitch decks with sharp, candid judgment. "
        "Return ONLY valid JSON with exactly this structure, no markdown:\n\n"
        '{"overall_score":<0-100>,"verdict":"<INVEST|MONITOR|PASS>",'
        '"one_liner":"<one sentence>","sections":{"problem":{"score":<0-100>,"comment":"<2 sentences>"},'
        '"solution":{"score":<0-100>,"comment":"<2 sentences>"},'
        '"market":{"score":<0-100>,"comment":"<2 sentences>"},'
        '"traction":{"score":<0-100>,"comment":"<2 sentences>"},'
        '"team":{"score":<0-100>,"comment":"<2 sentences>"},'
        '"financials":{"score":<0-100>,"comment":"<2 sentences>"}},'
        '"red_flags":["<1>","<2>","<3>","<4>","<5>"],'
        '"comparables":[{"name":"<co>","description":"<why>"},{"name":"<co>","description":"<why>"},'
        '{"name":"<co>","description":"<why>"}],'
        '"suggestions":["<1>","<2>","<3>","<4>","<5>"]}\n\n'
        "Rules: Be brutally honest. Most decks score 40-65. Reserve 80+ for exceptional. "
        "Verdict: INVEST>=72, MONITOR 50-71, PASS<50. Red flags must be specific to this deck."
    )
    user_msg = f"PITCH DECK:\n\n{pdf_text[:12000]}\n\nReturn the JSON analysis."
    resp = client.chat.completions.create(
        model=_model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_msg}],
        temperature=0.3, max_tokens=2000,
    )
    raw = resp.choices[0].message.content.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


def bar_color(s):
    return "#7ECFC8" if s >= 70 else ("#2A6B55" if s >= 45 else "#E09590")


def make_gauge_svg(score: int) -> str:
    r = 70
    cx = cy = 90
    circumference = 2 * 3.14159 * r
    arc = circumference * 0.75
    filled = arc * (score / 100)
    gap = arc - filled
    rotation = 135
    return f"""
<svg width="180" height="180" viewBox="0 0 180 180" xmlns="http://www.w3.org/2000/svg" class="ps-gauge-svg">
  <defs>
    <linearGradient id="gold-grad" x1="0" y1="0" x2="1" y2="0">
      <stop offset="0%" stop-color="#E8BA4A"/>
      <stop offset="100%" stop-color="#f0cc6a"/>
    </linearGradient>
    <linearGradient id="arc-grad" x1="0" y1="0" x2="1" y2="1">
      <stop offset="0%" stop-color="#7ECFC8"/>
      <stop offset="100%" stop-color="#E8BA4A"/>
    </linearGradient>
  </defs>
  <!-- Track -->
  <circle cx="{cx}" cy="{cy}" r="{r}" fill="none" stroke="rgba(255,255,255,0.05)" stroke-width="10"
    stroke-dasharray="{arc} {circumference - arc}"
    stroke-dashoffset="0"
    stroke-linecap="round"
    transform="rotate({rotation} {cx} {cy})"/>
  <!-- Fill (animated) -->
  <circle cx="{cx}" cy="{cy}" r="{r}" fill="none" stroke="url(#arc-grad)" stroke-width="10"
    stroke-dasharray="{arc} {circumference}"
    stroke-dashoffset="{arc}"
    stroke-linecap="round"
    transform="rotate({rotation} {cx} {cy})">
    <animate attributeName="stroke-dashoffset"
      from="{arc}" to="{gap}"
      dur="1.5s" begin="0.1s" fill="freeze"
      calcMode="spline" keyTimes="0;1" keySplines="0.32,0.72,0,1"/>
  </circle>
  <!-- Score -->
  <text x="{cx}" y="{cy - 6}" text-anchor="middle" dominant-baseline="middle"
    font-size="36" font-weight="800" fill="url(#gold-grad)" font-family="'Fira Code',monospace">{score}</text>
  <text x="{cx}" y="{cy + 22}" text-anchor="middle"
    font-size="9" fill="#475569" font-family="Inter,sans-serif" letter-spacing="2" text-transform="uppercase">VC SCORE</text>
</svg>"""


def build_markdown_report(result: dict, filename: str) -> str:
    r = result
    score = r["overall_score"]
    verdict = r["verdict"]
    lines = [
        f"# PitchScan Report — {filename}",
        "",
        f"**Overall Score:** {score} / 100  ",
        f"**Verdict:** {verdict}  ",
        f"**Summary:** {r.get('one_liner', '')}",
        "",
        "---",
        "",
        "## Section Breakdown",
        "",
    ]
    for key, val in r.get("sections", {}).items():
        lines.append(f"### {key.capitalize()} — {val['score']}/100")
        lines.append(val["comment"])
        lines.append("")

    lines += ["---", "", "## Red Flags", ""]
    for flag in r.get("red_flags", []):
        lines.append(f"- ⚠ {flag}")

    lines += ["", "---", "", "## Comparable Companies", ""]
    for c in r.get("comparables", []):
        lines.append(f"**{c['name']}** — {c['description']}")
        lines.append("")

    lines += ["---", "", "## Improvement Suggestions", ""]
    for s in r.get("suggestions", []):
        lines.append(f"- → {s}")

    lines += ["", "---", "", "*Generated by PitchScan · Powered by DeepSeek AI*"]
    return "\n".join(lines)


def build_docx_report(result: dict, filename: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"

    title = doc.add_heading(f"PitchScan Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"File: {filename}")
    doc.add_paragraph(f"Overall Score: {result['overall_score']} / 100")
    doc.add_paragraph(f"Verdict: {result['verdict']}")
    doc.add_paragraph(f"Summary: {result.get('one_liner','')}")
    doc.add_paragraph("")

    doc.add_heading("Section Breakdown", 1)
    for key, val in result.get("sections", {}).items():
        p = doc.add_paragraph()
        run = p.add_run(f"{key.capitalize()} — {val['score']}/100")
        run.bold = True
        doc.add_paragraph(val["comment"])

    doc.add_heading("Red Flags", 1)
    for flag in result.get("red_flags", []):
        doc.add_paragraph(f"⚠ {flag}", style="List Bullet")

    doc.add_heading("Comparable Companies", 1)
    for c in result.get("comparables", []):
        p = doc.add_paragraph()
        p.add_run(c["name"] + " — ").bold = True
        p.add_run(c["description"])

    doc.add_heading("Improvement Suggestions", 1)
    for s in result.get("suggestions", []):
        doc.add_paragraph(f"→ {s}", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_results(result: dict, filename: str):
    import streamlit.components.v1 as _components
    score = result["overall_score"]
    verdict = result["verdict"]
    one_liner = result.get("one_liner", "")

    # Top row: gauge + sections
    col_g, col_s = st.columns([1, 2])
    with col_g:
        st.markdown(
            f'<div class="ps-gauge-wrap">'
            f'{make_gauge_svg(score)}'
            f'<div><span class="verdict-chip verdict-{verdict}">'
            f'<span class="verdict-dot"></span>{verdict}'
            f'</span></div>'
            f'<div class="ps-oneliner">{one_liner}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    with col_s:
        st.markdown('<div class="ps-sh">Section Breakdown</div>', unsafe_allow_html=True)
        bars = ""
        for key, val in result.get("sections", {}).items():
            s = val["score"]
            color = bar_color(s)
            bars += (
                f'<div class="ps-bar-row">'
                f'<span class="ps-bar-name">{key.capitalize()}</span>'
                f'<div class="ps-bar-bg"><div class="ps-bar-fill" style="width:{s}%;background:{color};"></div></div>'
                f'<span class="ps-bar-score">{s}</span>'
                f'</div>'
            )
        st.markdown(f'<div class="ps-bars">{bars}</div>', unsafe_allow_html=True)
        with st.expander("View section comments"):
            for key, val in result.get("sections", {}).items():
                st.markdown(f"**{key.capitalize()}** — {val['comment']}")

    st.markdown("<hr>", unsafe_allow_html=True)

    # Flags + Suggestions
    col_f, col_sg = st.columns(2)
    with col_f:
        st.markdown('<div class="ps-sh">Red Flags</div>', unsafe_allow_html=True)
        html = "".join(f'<div class="ps-flag">&#9888;&nbsp; {f}</div>' for f in result.get("red_flags", []))
        st.markdown(html, unsafe_allow_html=True)
    with col_sg:
        st.markdown('<div class="ps-sh">Improvement Suggestions</div>', unsafe_allow_html=True)
        html = "".join(f'<div class="ps-sug">&#8594;&nbsp; {s}</div>' for s in result.get("suggestions", []))
        st.markdown(html, unsafe_allow_html=True)

    # Comparables
    st.markdown('<div class="ps-sh">Comparable Companies</div>', unsafe_allow_html=True)
    comp_cols = st.columns(3)
    for i, c in enumerate(result.get("comparables", [])):
        with comp_cols[i % 3]:
            st.markdown(
                f'<div class="ps-comp">'
                f'<div class="ps-comp-name">{c["name"]}</div>'
                f'<div class="ps-comp-desc">{c["description"]}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

    # Downloads
    st.markdown('<div class="ps-sh">Download Report</div>', unsafe_allow_html=True)
    base_name = filename.replace(".pdf", "")
    dcol1, dcol2, dcol3, _ = st.columns([1, 1, 1, 2])
    with dcol1:
        st.download_button(
            "JSON",
            data=json.dumps(result, indent=2),
            file_name=f"pitchscan_{base_name}.json",
            mime="application/json",
        )
    with dcol2:
        st.download_button(
            "Markdown",
            data=build_markdown_report(result, filename),
            file_name=f"pitchscan_{base_name}.md",
            mime="text/markdown",
        )
    with dcol3:
        try:
            docx_bytes = build_docx_report(result, filename)
            st.download_button(
                "Word (DOCX)",
                data=docx_bytes,
                file_name=f"pitchscan_{base_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception:
            pass

    # Feedback
    st.markdown('<div class="ps-sh">Was this analysis useful?</div>', unsafe_allow_html=True)
    _components.html("""
<style>
.fb-row{display:flex;align-items:center;gap:12px;margin:4px 0 8px;}
.fb-btn{
  background:rgba(12,42,56,0.5);border:1px solid rgba(126,207,200,0.12);
  border-radius:8px;padding:8px 20px;cursor:pointer;font-size:1.1rem;
  transition:all 0.2s;color:#e8f4f2;font-family:system-ui;
}
.fb-btn:hover{border-color:rgba(232,186,74,0.45);transform:scale(1.08);}
.fb-btn.selected{border-color:#E8BA4A;background:rgba(232,186,74,0.12);transform:scale(1.08);}
.fb-msg{font-size:0.78rem;color:#5a7a82;font-family:system-ui;margin-left:4px;}
</style>
<div class="fb-row">
  <button class="fb-btn" id="fb-up" onclick="vote(1)">👍</button>
  <button class="fb-btn" id="fb-dn" onclick="vote(-1)">👎</button>
  <span class="fb-msg" id="fb-msg"></span>
</div>
<script>
function vote(v){
  document.getElementById("fb-up").classList.remove("selected");
  document.getElementById("fb-dn").classList.remove("selected");
  if(v>0){
    document.getElementById("fb-up").classList.add("selected");
    document.getElementById("fb-msg").textContent = "Thanks! Glad it helped.";
  } else {
    document.getElementById("fb-dn").classList.add("selected");
    document.getElementById("fb-msg").textContent = "Got it. We'll keep improving.";
  }
}
</script>
""", height=70, scrolling=False)

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("Scan Another Deck"):
        for k in ("result", "last_file"):
            st.session_state.pop(k, None)
        st.rerun()


def _is_builtin_key(key: str) -> bool:
    """True when the key came from st.secrets (not user-entered)."""
    try:
        return key == st.secrets.get("DEEPSEEK_API_KEY", "")
    except Exception:
        return False


def main():
    import streamlit.components.v1 as _components
    _components.html(f"""
<script>
(function(n,o,v,u,s,ai){{
  if(n[s])return;n[s]=Object.assign(n[s]||{{}},{{initialize:function(opts){{n[s].q=(n[s].q||[]).concat([["initialize",opts]]);}} }});
  n[s].initialize({{apiKey:"{NOVUS_APP_ID}"}});
  ai=o.createElement('script');ai.async=true;
  ai.src='https://cdn.novus.pendo.io/agent/static/{NOVUS_APP_ID}/novus.js';
  o.head.appendChild(ai);
}})(window,document,'novus','novus','novus');
</script>""", height=1, scrolling=False)
    _components.html("""
<script>
(function(apiKey){
    (function(p,e,n,d,o){var v,w,x,y,z;o=p[d]=p[d]||{};o._q=o._q||[];
    v=['initialize','identify','updateOptions','pageLoad','track', 'trackAgent'];for(w=0,x=v.length;w<x;++w)(function(m){
    o[m]=o[m]||function(){o._q[m===v[0]?'unshift':'push']([m].concat([].slice.call(arguments,0)));};})(v[w]);
    y=e.createElement(n);y.async=!0;y.src='https://cdn.pendo.io/agent/static/'+apiKey+'/pendo.js';
    z=e.getElementsByTagName(n)[0];z.parentNode.insertBefore(y,z);})(window,document,'script','pendo');
})('f2950f7e-4ee3-4154-81f4-c3f25610a0fc');
pendo.initialize({ visitor: { id: '' } });
</script>
""", height=0, scrolling=False)

    st.markdown(CSS, unsafe_allow_html=True)

    # Sidebar: API key
    _has_builtin = False
    try:
        _secret = st.secrets.get("DEEPSEEK_API_KEY", "")
        if _secret:
            api_key = _secret
            _has_builtin = True
    except Exception:
        pass

    # Provider config map
    _PROVIDERS = {
        "DeepSeek": {
            "base_url": "https://api.deepseek.com",
            "model": "deepseek-chat",
            "placeholder": "sk-...",
            "link": "https://platform.deepseek.com",
            "link_label": "Get free DeepSeek key →",
            "note": "~$0.004 per analysis",
        },
        "OpenAI (GPT-4o)": {
            "base_url": "https://api.openai.com/v1",
            "model": "gpt-4o",
            "placeholder": "sk-...",
            "link": "https://platform.openai.com/api-keys",
            "link_label": "Get OpenAI key →",
            "note": "~$0.03 per analysis",
        },
        "Groq (free tier)": {
            "base_url": "https://api.groq.com/openai/v1",
            "model": "llama-3.3-70b-versatile",
            "placeholder": "gsk_...",
            "link": "https://console.groq.com/keys",
            "link_label": "Get free Groq key →",
            "note": "Free tier available",
        },
    }

    with st.sidebar:
        st.markdown("### AI Provider")
        _provider_name = st.selectbox(
            "Choose provider",
            list(_PROVIDERS.keys()),
            label_visibility="collapsed",
        )
        _prov = _PROVIDERS[_provider_name]

        st.markdown("### API Key")
        _user_key = st.text_input(
            f"{_provider_name} API key",
            type="password",
            placeholder=_prov["placeholder"],
            label_visibility="collapsed",
        )
        if _user_key:
            api_key = _user_key
            st.success("🔒 Your key is used only for this analysis and is never stored.")
        elif _has_builtin and _provider_name == "DeepSeek":
            _rem = daily_remaining()
            st.caption(f"Using shared key · **{_rem} free analyses left today**")
        st.caption(f"[{_prov['link_label']}]({_prov['link']}) · {_prov['note']}")

        # Store provider config in session for use in analyze_pitch
        st.session_state["_provider"] = _prov

    # Hero with counter
    _deck_count = get_deck_count()
    _remaining = daily_remaining()
    _using_own_key = bool(api_key and not _is_builtin_key(api_key))
    if _using_own_key:
        _badge_extra = "Your API key · Unlimited"
    else:
        _badge_extra = f"{_remaining} free analyses left today"
    st.markdown(
        f'<div class="ps-hero">'
        f'<div class="ps-hero-badge"><span class="ps-hero-dot"></span>AI-Powered VC Analysis &nbsp;·&nbsp; {_deck_count} decks analyzed &nbsp;·&nbsp; {_badge_extra}</div>'
        f'<h1>Know if your pitch will<br><span>get funded</span></h1>'
        f'<p>Upload your pitch deck and get an instant, brutally honest VC-grade review — '
        f'score, red flags, comparable companies, and concrete improvements.</p>'
        f'</div>',
        unsafe_allow_html=True,
    )

    # Survey context form
    with st.expander("📋 Add context for a better analysis (optional)", expanded=False):
        sc1, sc2 = st.columns(2)
        with sc1:
            survey_stage = st.selectbox(
                "Fundraising stage",
                ["", "Pre-seed", "Seed", "Series A", "Series B+"],
                index=0,
            )
            survey_sector = st.selectbox(
                "Sector",
                ["", "B2B SaaS", "Consumer", "Fintech", "Healthtech", "Deep Tech", "Marketplace", "Other"],
                index=0,
            )
            survey_revenue = st.selectbox(
                "Revenue status",
                ["", "Pre-revenue", "<$10k MRR", "$10k–$100k MRR", "$100k+ MRR"],
                index=0,
            )
        with sc2:
            survey_investor = st.selectbox(
                "Target investor type",
                ["", "YC", "Angels", "VC Tier 1", "VC Regional", "Strategic"],
                index=0,
            )
            survey_persona = st.selectbox(
                "Evaluate as…",
                ["", "YC", "a16z", "Sequoia", "Benchmark"],
                index=0,
                help="Changes the investor lens used for the analysis",
            )
            survey_concerns = st.multiselect(
                "Areas you're most worried about",
                ["Market size", "Traction story", "Team credibility", "Business model", "Competition", "Storytelling"],
            )
    survey_context = {
        "stage": survey_stage,
        "sector": survey_sector,
        "revenue": survey_revenue,
        "investor_type": survey_investor,
        "persona": survey_persona,
        "concerns": survey_concerns,
    }

    # Upload
    uploaded_file = st.file_uploader(
        "Drop your pitch deck PDF here",
        type=["pdf"],
        label_visibility="collapsed",
    )
    st.markdown(
        '<div class="ps-upload-hint">Supports text-based and image-only PDFs &middot; '
        'Analysis takes a few minutes</div>',
        unsafe_allow_html=True,
    )

    if not uploaded_file:
        st.markdown("<br>", unsafe_allow_html=True)
        c1, c2, c3, c4 = st.columns(4)
        for col, icon, label, val in [
            (c1, "🎯", "VC Score", "0 – 100"),
            (c2, "🚩", "Red Flags", "Identified"),
            (c3, "🔍", "Comparables", "3 Companies"),
            (c4, "💡", "Suggestions", "5 Actions"),
        ]:
            col.markdown(
                f'<div style="background:rgba(255,255,255,0.02);border:1px solid rgba(255,255,255,0.05);'
                f'border-radius:12px;padding:1.2rem;text-align:center;">'
                f'<div style="font-size:1.8rem;margin-bottom:0.4rem;">{icon}</div>'
                f'<div style="color:#475569;font-size:0.72rem;text-transform:uppercase;letter-spacing:1px;">{label}</div>'
                f'<div style="color:#e2e8f0;font-weight:700;font-size:1rem;margin-top:0.3rem;">{val}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
        return

    if not api_key:
        st.warning("Add your DeepSeek API key in the sidebar to run the analysis.")
        return

    # Daily cap check (only applies when using built-in key)
    _builtin = _is_builtin_key(api_key)
    if _builtin and daily_cap_reached():
        st.markdown(
            '<div style="background:rgba(232,186,74,0.08);border:1px solid rgba(232,186,74,0.25);'
            'border-radius:12px;padding:1.2rem 1.4rem;margin:1rem 0;">' 
            '<div style="color:#E8BA4A;font-weight:700;margin-bottom:0.4rem;">'
            '⏳ Free daily quota reached (20 analyses/day)</div>'
            '<div style="color:#5a7a82;font-size:0.85rem;line-height:1.6;">'
            'Add your own <a href="https://platform.deepseek.com" target="_blank" '
            'style="color:#7ECFC8;">DeepSeek API key</a> in the sidebar to keep going — '
            'it&#39;s free to sign up and costs ~$0.004 per analysis. '
            'Or check back tomorrow when the quota resets.</div></div>',
            unsafe_allow_html=True,
        )
        return

    cache_key = uploaded_file.name + str(uploaded_file.size) + json.dumps(survey_context, sort_keys=True)

    if "result" not in st.session_state or st.session_state.get("last_file") != cache_key:
        render_loading_tips()

        with st.spinner("Reading pitch deck..."):
            try:
                pdf_bytes = uploaded_file.read()
                pitch_text = extract_text_from_pdf(pdf_bytes)
            except Exception as e:
                st.error(f"Could not read PDF: {e}")
                return

        if not pitch_text.strip():
            st.error("Could not extract any text from this PDF. Try a different file.")
            return

        with st.spinner("Analyzing with DeepSeek — this takes a few minutes..."):
            try:
                result = analyze_pitch(pitch_text, api_key, context=survey_context)
                increment_deck_count()
                st.session_state["result"] = result
                st.session_state["last_file"] = cache_key
            except json.JSONDecodeError as e:
                st.error(f"Model returned invalid JSON. Please try again.")
                return
            except Exception as e:
                st.error(f"Analysis failed: {e}")
                return
        st.rerun()

    render_results(st.session_state["result"], uploaded_file.name)

    st.markdown(
        '<div style="text-align:center;margin-top:2rem;color:#1e293b;font-size:0.72rem;">'
        'PitchScan &middot; World Product Day Hackathon 2026 &middot; Powered by DeepSeek AI'
        '</div>',
        unsafe_allow_html=True,
    )


main()
